#!/usr/bin/env python3
"""Genera los 5 documentos legales de NuncaCierro en formato .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime, date

# ─── CONFIG ───────────────────────────────────────────────────────
COMPANY = "NuncaCierro"
SLOGAN = "NuncaCierro"
EMAIL = "soporte@nuncacierro.com"
WHATSAPP = "+57 3219615338"
DOMAIN = "nuncacierro.com"
CITY = "Bogotá, Colombia"
LEGAL_REP = "Nicolas Valenzuela"
LEGAL_REP_CC = "CC 1.020.825.162"
LOGO_PATH = os.path.abspath(
    r"D:\Documentos\Repositories\nunca-cierro\nc-dashboard\public\Logonobg.png"
)
OUT_DIR = os.path.abspath(r"D:\Documentos\Repositories\nunca-cierro\docs\legal")

# Colores corporativos
DARK_BLUE = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
MEDIUM_BLUE = RGBColor(0x1E, 0x3A, 0x5F)  # #1E3A5F
ACCENT = RGBColor(0x25, 0x6E, 0xB5)       # #256EB5
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
VERY_LIGHT = RGBColor(0xF1, 0xF5, 0xF9)

TODAY = date.today().strftime("%d de %B de %Y")

# ─── HELPERS ──────────────────────────────────────────────────────

def set_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_header(doc, show_logo=True):
    """Agrega encabezado corporativo con logo y datos."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Tabla invisible para layout: logo | datos
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Estilo sin bordes
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
        tblPr.append(borders)

        # Celda 1: Logo
        cell_logo = table.cell(0, 0)
        cell_logo.width = Inches(1.2)
        cell_logo_p = cell_logo.paragraphs[0]
        cell_logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if show_logo and os.path.exists(LOGO_PATH):
            run = cell_logo_p.add_run()
            run.add_picture(LOGO_PATH, width=Inches(0.7))

        # Celda 2: Datos empresa
        cell_info = table.cell(0, 1)
        cell_info.width = Inches(5.3)
        cp = cell_info.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = cp.add_run(f"{COMPANY}\n")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"
        run2 = cp.add_run(f"{LEGAL_REP} · {LEGAL_REP_CC}\n{EMAIL} | {WHATSAPP}\n{domain_url()}")
        run2.font.size = Pt(8)
        run2.font.color.rgb = LIGHT_GRAY
        run2.font.name = "Calibri"

        # Línea separadora
        p_space = header.add_paragraph()
        p_space.space_before = Pt(2)
        p_space.space_after = Pt(0)
        run_line = p_space.add_run()
        run_line.font.size = Pt(1)
        # Add bottom border to paragraph
        pPr = p_space._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="256EB5"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)


def add_footer(doc):
    """Agrega pie de página con número de página."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = doc.styles['Normal']
        
        # Línea superior
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="1" w:color="256EB5"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)

        run = p.add_run(f"{COMPANY} · ")
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GRAY
        run.font.name = "Calibri"

        # Número de página
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2 = p.add_run()
        run2._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3 = p.add_run()
        run3._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4 = p.add_run()
        run4._r.append(fldChar2)
        for r in [run2, run3, run4]:
            r.font.size = Pt(8)
            r.font.color.rgb = LIGHT_GRAY
            r.font.name = "Calibri"


def domain_url():
    return f"https://{DOMAIN}"


def add_title(doc, text, level=0):
    """Agrega un título estilizado."""
    if level == 0:
        # Título principal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.space_before = Pt(6)
        p.space_after = Pt(12)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"
        # Línea decorativa
        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(12)
        run2 = p2.add_run()
        run2.font.size = Pt(1)
        pPr = p2._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="256EB5"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
        return p
    elif level == 1:
        p = doc.add_paragraph()
        p.space_before = Pt(14)
        p.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = MEDIUM_BLUE
        run.font.name = "Calibri"
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p.space_before = Pt(10)
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_BLUE
        run.font.name = "Calibri"
        return p


def add_body(doc, text, bold=False, italic=False, size=10.5, color=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.space_before = Pt(2)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK_GRAY
    run.font.name = "Calibri"
    return p


def add_field_row(doc, label, value=""):
    """Campo label: value en línea."""
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = DARK_BLUE
    run_label.font.name = "Calibri"
    run_value = p.add_run(value)
    run_value.font.size = Pt(10)
    run_value.font.color.rgb = DARK_GRAY
    run_value.font.name = "Calibri"
    # Subrayado para campos editables
    run_value.underline = True
    return p


def add_blank_field(doc, label, width_cm=8):
    """Campo label + línea para rellenar."""
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(2)
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = DARK_BLUE
    run_label.font.name = "Calibri"
    # Línea subrayada
    run_line = p.add_run("_" * 80)
    run_line.font.size = Pt(10)
    run_line.font.color.rgb = LIGHT_GRAY
    run_line.font.name = "Calibri"
    return p


def add_signature_line(doc, label):
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("_" * 45)
    run.font.color.rgb = LIGHT_GRAY
    run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.space_before = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(label)
    run2.bold = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK_BLUE
    run2.font.name = "Calibri"
    return p, p2


def add_separator(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    run = p.add_run()
    run.font.size = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
    run = p.add_run(f"  {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Calibri"
    return p


def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold or header
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"
        if header:
            run.font.color.rgb = WHITE
            # Fondo azul
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A5F"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        else:
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    return row


# ─── DOCUMENT 1: COTIZACIÓN ──────────────────────────────────────

def create_cotizacion():
    doc = Document()
    set_margins(doc)
    add_header(doc)
    add_footer(doc)

    add_title(doc, "Cotización")

    # Número y fecha
    add_field_row(doc, "Cotización No", "NC-____-___")
    add_field_row(doc, "Fecha", TODAY)

    add_separator(doc)
    add_title(doc, "Datos del Cliente", level=2)
    add_blank_field(doc, "Cliente / Empresa")
    add_blank_field(doc, "Correo electrónico")
    add_blank_field(doc, "Teléfono / WhatsApp")

    add_separator(doc)
    add_title(doc, "Descripción del Proyecto", level=1)
    add_body(doc, "Describir brevemente el proyecto y sus objetivos:", italic=True, size=10)
    add_blank_field(doc, "Descripción")
    add_blank_field(doc, "Objetivos")
    add_blank_field(doc, "Alcance")

    add_title(doc, "Tecnologías", level=2)
    add_body(doc, "Next.js · React · Tailwind CSS · TypeScript · Vercel", size=10)

    add_title(doc, "Tiempo Estimado", level=2)
    add_blank_field(doc, "Tiempo de desarrollo")

    add_separator(doc)
    add_title(doc, "Valores y Forma de Pago", level=1)

    # Tabla de valores
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    add_table_row(table, ["Concepto", "Valor"], header=True)
    add_table_row(table, ["Desarrollo del proyecto", "$ _______ COP"])
    add_table_row(table, ["Dominio + Hosting (1 año)", "$ _______ COP"])
    add_table_row(table, ["Total", "$ _______ COP"], bold=True)

    add_separator(doc)
    add_blank_field(doc, "Forma de pago")
    add_blank_field(doc, "Condiciones adicionales")

    add_separator(doc)
    add_title(doc, "Firma", level=2)
    add_signature_line(doc, f"{LEGAL_REP} — {LEGAL_REP_CC}")
    add_signature_line(doc, "Cliente")

    doc.save(os.path.join(OUT_DIR, "01_Cotizacion.docx"))
    print("[OK] 01_Cotizacion.docx")


# ─── DOCUMENT 2: CONTRATO ────────────────────────────────────────

def create_contrato():
    doc = Document()
    set_margins(doc)
    add_header(doc)
    add_footer(doc)

    add_title(doc, "Contrato de Prestación de Servicios")

    add_body(doc, f"Entre {COMPANY}, representada por {LEGAL_REP} ({LEGAL_REP_CC}), en adelante \"el desarrollador\", y el cliente que se identifica al final de este documento, se celebra el presente contrato de prestación de servicios profesionales de desarrollo web, el cual se rige por las siguientes cláusulas:", size=10)

    add_separator(doc)
    clausulas = [
        ("PRIMERA — OBJETO",
         "El desarrollador se obliga a diseñar, desarrollar y publicar un sitio web moderno y responsivo para el cliente, utilizando las tecnologías acordadas (Next.js, React, Tailwind CSS) y desplegándolo en Vercel, de acuerdo con las especificaciones y alcance definidos en la cotización aceptada por el cliente."),
        ("SEGUNDA — ALCANCE",
         "El alcance del proyecto incluye: (a) Diseño y desarrollo del número de secciones acordadas, (b) Adaptación responsiva (celular, tablet, escritorio), (c) Integración de formulario de contacto y enlace WhatsApp, (d) Despliegue en producción. Cualquier funcionalidad no listada en la cotización será considerada como cambio adicional."),
        ("TERCERA — VALOR Y FORMA DE PAGO",
         "El valor total del proyecto es el indicado en la cotización firmada. El pago se realizará así: 50% al inicio del proyecto y 50% contra entrega, salvo que se acuerde otra forma de pago por escrito. Los pagos se reciben en pesos colombianos (COP) por transferencia bancaria o efectivo."),
        ("CUARTA — RESPONSABILIDADES DEL DESARROLLADOR",
         "Entregar el sitio web funcional en el tiempo acordado. Garantizar que el sitio sea responsivo y funcione correctamente. Realizar hasta 2 rondas de correcciones sobre el contenido suministrado por el cliente. Entregar acceso al repositorio y al panel de administración cuando aplique. Prestar soporte técnico por 15 días calendario después de la entrega para corrección de errores técnicos."),
        ("QUINTA — RESPONSABILIDADES DEL CLIENTE",
         "Entregar toda la información necesaria (textos, imágenes, logos, colores) antes de iniciar el desarrollo. Revisar y dar aprobación por escrito en cada etapa. Realizar los pagos en las fechas acordadas. Proveer acceso a redes sociales, dominios o servicios de terceros cuando sea necesario."),
        ("SEXTA — TIEMPOS DE ENTREGA",
         "El tiempo estimado de entrega es el indicado en la cotización, contado a partir del recibo del anticipo y de la información completa del cliente. Los retrasos por parte del cliente en la entrega de información corren el cronograma automáticamente."),
        ("SÉPTIMA — REVISIONES Y CAMBIOS",
         "Se incluyen 2 rondas de revisión y corrección sobre el contenido existente. Los cambios adicionales o nuevas funcionalidades no contempladas en el alcance inicial se cotizarán por separado y requerirán aprobación del cliente antes de ser ejecutados."),
        ("OCTAVA — PROPIEDAD INTELECTUAL",
         "Una vez cancelado la totalidad del valor del proyecto, el cliente recibe la propiedad intelectual completa del código fuente y los activos digitales creados específicamente para su sitio web. El desarrollador se reserva el derecho de incluir el proyecto en su portafolio profesional, salvo acuerdo de confidencialidad firmado."),
        ("NOVENA — ENTREGA DEL PROYECTO",
         "Se entenderá entregado cuando: (a) el sitio esté publicado y accesible en el dominio acordado, (b) se haya firmado el acta de entrega, y (c) se haya recibido el pago total del saldo pendiente."),
        ("DÉCIMA — TERMINACIÓN",
         "Cualquiera de las partes puede terminar el contrato por incumplimiento grave, dando aviso por escrito con 5 días hábiles de antelación. En caso de terminación anticipada por parte del cliente, los pagos ya realizados no serán reembolsables y se facturará el trabajo ejecutado hasta la fecha de terminación."),
    ]

    for title, body in clausulas:
        add_title(doc, title, level=2)
        add_body(doc, body, size=10)

    add_separator(doc)
    add_title(doc, "Firmas", level=1)
    add_body(doc, f"Dado en {CITY}, a los {TODAY}.", size=10, italic=True)

    add_signature_line(doc, f"{COMPANY}")
    add_body(doc, f"{LEGAL_REP} — {LEGAL_REP_CC}", size=10, italic=True)
    add_signature_line(doc, "Cliente — Nombre / Empresa")
    add_blank_field(doc, "CC / NIT")

    doc.save(os.path.join(OUT_DIR, "02_Contrato_Servicios.docx"))
    print("[OK] 02_Contrato_Servicios.docx")


# ─── DOCUMENT 3: RECIBO DE CAJA ──────────────────────────────────

def create_recibo_caja():
    doc = Document()
    set_margins(doc, top=2, bottom=2, left=2.5, right=2.5)
    add_header(doc)
    add_footer(doc)

    add_title(doc, "Recibo de Caja")

    # Número consecutivo
    add_field_row(doc, "Recibo No", "____-___")
    add_field_row(doc, "Fecha", TODAY)

    add_separator(doc)

    # Datos del recibí
    add_title(doc, "Recibí de", level=1)
    add_blank_field(doc, "Nombre / Razón Social")
    add_blank_field(doc, "Identificación (CC / NIT)")
    add_blank_field(doc, "Teléfono / Correo")

    add_separator(doc)

    # Valor
    add_title(doc, "Información del Pago", level=1)
    add_blank_field(doc, "Valor recibido ($)")
    add_blank_field(doc, "Valor en letras")
    add_blank_field(doc, "Concepto")
    add_blank_field(doc, "Forma de pago (Efectivo / Transferencia / Nequi / Daviplata)")
    add_blank_field(doc, "Observaciones")

    add_separator(doc)
    add_title(doc, "Firma", level=2)
    add_body(doc, "Recibí conforme:", size=10, italic=True)
    add_signature_line(doc, f"{LEGAL_REP} — {LEGAL_REP_CC}")

    doc.save(os.path.join(OUT_DIR, "03_Recibo_Caja.docx"))
    print("[OK] 03_Recibo_Caja.docx")


# ─── DOCUMENT 4: ACTA DE ENTREGA ─────────────────────────────────

def create_acta_entrega():
    doc = Document()
    set_margins(doc)
    add_header(doc)
    add_footer(doc)

    add_title(doc, "Acta de Entrega")

    add_body(doc, "Por medio de la presente, se deja constancia de la entrega formal del proyecto de desarrollo web descrito a continuación:", size=10, italic=True)

    add_separator(doc)
    add_blank_field(doc, "Cliente")
    add_blank_field(doc, "Empresa / Razón Social")
    add_blank_field(doc, "Proyecto / Sitio web")
    add_field_row(doc, "Fecha de entrega", TODAY)

    add_separator(doc)
    add_title(doc, "Descripción de lo Entregado", level=1)
    add_body(doc, "Se entrega sitio web funcional, responsivo y optimizado, desarrollado con las tecnologías acordadas:", size=10)
    add_bullet(doc, "Sitio web completo con las secciones acordadas")
    add_bullet(doc, "Diseño responsivo (celular, tablet, escritorio)")
    add_bullet(doc, "Formulario de contacto operativo")
    add_bullet(doc, "Integración de enlace WhatsApp")
    add_bullet(doc, "Contenido y textos proporcionados por el cliente")

    add_separator(doc)
    add_title(doc, "Accesos y Recursos Entregados", level=1)
    add_blank_field(doc, "Repositorio (GitHub/Bitbucket)")
    add_blank_field(doc, "Dominio")
    add_blank_field(doc, "URL del sitio en producción")
    add_blank_field(doc, "Dashboard Vercel / Hosting")
    add_blank_field(doc, "Otros accesos (Google Analytics, etc.)")

    add_separator(doc)
    add_title(doc, "Observaciones", level=1)
    add_blank_field(doc, "Observaciones")

    add_separator(doc)
    add_title(doc, "Conformidad", level=1)
    add_body(doc, "El cliente declara recibir el proyecto a satisfacción y conforme con lo acordado.", size=10, italic=True)

    add_signature_line(doc, "Cliente — Nombre y firma")
    add_signature_line(doc, "CC / NIT del cliente")
    add_signature_line(doc, f"{COMPANY}")
    add_body(doc, f"{LEGAL_REP} — {LEGAL_REP_CC}", size=10, italic=True)

    doc.save(os.path.join(OUT_DIR, "04_Acta_Entrega.docx"))
    print("[OK] 04_Acta_Entrega.docx")


# ─── DOCUMENT 5: PROPUESTA COMERCIAL ─────────────────────────────

def create_propuesta():
    doc = Document()
    set_margins(doc, top=3, bottom=2.5, left=2.5, right=2.5)

    # ── PORTADA ──
    # Espacio superior
    for _ in range(4):
        p = doc.add_paragraph()
        p.space_after = Pt(0)

    # Logo centrado
    if os.path.exists(LOGO_PATH):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.5))

    # Título principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.space_before = Pt(20)
    run = p_title.add_run("Propuesta Comercial")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    # Línea decorativa
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("_________________________________")
    run_line.font.color.rgb = ACCENT
    run_line.font.size = Pt(12)

    # Subtítulo
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.space_before = Pt(8)
    run = p_sub.add_run("Sitio web profesional con WhatsApp\npara tu negocio")
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = "Calibri"

    # Fecha
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.space_before = Pt(30)
    run = p_date.add_run(TODAY)
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_GRAY

    # Cliente
    p_client = doc.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_client.add_run("Presentado a: ____________________________")
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_GRAY

    # Salto de página
    doc.add_page_break()

    # ── CONTENIDO INTERNO con header/footer ──
    # Ahora agregamos header y footer a todas las secciones de contenido
    add_header(doc)
    add_footer(doc)

    add_title(doc, "1. Presentación")
    add_body(doc, f"{COMPANY} es una empresa colombiana especializada en el desarrollo de sitios web modernos y profesionales. Creamos presencia digital para negocios que quieren vender más y atender mejor a sus clientes, combinando diseño web de alta calidad con la potencia de WhatsApp.")
    add_body(doc, f"Nuestro enfoque es claro: sitios rápidos, modernos y que generen resultados reales. No hacemos páginas genéricas — diseñamos cada proyecto pensando en el negocio del cliente.")

    add_title(doc, "2. Servicios")
    servicios = [
        "Sitio web profesional one-page o multi-sección",
        "Diseño responsivo (celular, tablet, escritorio)",
        "Integración de WhatsApp para atención al cliente",
        "Formulario de contacto automatizado",
        "Optimización básica para buscadores (SEO)",
        " Hosting y dominio gestionados (opcional)",
    ]
    for s in servicios:
        add_bullet(doc, s)

    add_title(doc, "3. Tecnologías")
    add_body(doc, "Trabajamos con tecnología de última generación para garantizar velocidad, seguridad y una experiencia de usuario excepcional:")
    techs = [
        ("Next.js 16", "El framework de React más avanzado"),
        ("React 19", "Interfaces dinámicas y reactivas"),
        ("Tailwind CSS v4", "Diseño moderno y responsivo"),
        ("TypeScript", "Código robusto y mantenible"),
        ("Vercel", "Hosting en la nube con despliegue automático"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    add_table_row(table, ["Tecnología", "Beneficio"], header=True)
    for tech, desc in techs:
        add_table_row(table, [tech, desc])

    add_title(doc, "4. Proceso de Trabajo")
    pasos = [
        ("1. Diagnóstico", "Entendemos tu negocio, tu público y lo que necesitas lograr."),
        ("2. Propuesta", "Te compartimos diseño, alcance, tiempos y presupuesto."),
        ("3. Desarrollo", "Construimos tu sitio con entregas periódicas para tu revisión."),
        ("4. Correcciones", "Hasta 2 rondas de ajustes sobre el contenido."),
        ("5. Publicación", "Desplegamos tu sitio y te entregamos accesos."),
        ("6. Soporte", "Te acompañamos post-entrega para asegurar que todo funcione."),
    ]
    for paso, desc in pasos:
        add_title(doc, paso, level=2)
        add_body(doc, desc, size=10)

    add_title(doc, "5. Cronograma Estimado")
    add_body(doc, "El tiempo total depende de la complejidad del proyecto y de la rapidez con que recibamos la información del cliente:", size=10)
    add_bullet(doc, "Plan Básico: 3 – 5 días hábiles")
    add_bullet(doc, "Plan Profesional: 5 – 10 días hábiles")
    add_bullet(doc, "Plan Premium: 10 – 15 días hábiles")

    add_title(doc, "6. Beneficios de Trabajar con Nosotros")
    beneficios = [
        "Sitio web moderno que refleja la seriedad de tu negocio",
        "Cargado rápido y optimizado para celular",
        "WhatsApp integrado para recibir clientes 24/7",
        "Sin complicaciones técnicas — nosotros nos encargamos de todo",
        "Soporte directo por WhatsApp, sin vueltas",
    ]
    for b in beneficios:
        add_bullet(doc, b)

    add_title(doc, "7. Inversión")
    add_body(doc, "Los planes y precios se detallan en la cotización adjunta. Trabajamos con esquemas de pago flexibles para adaptarnos a tu flujo de caja.", size=10)

    # Tabla de planes
    table2 = doc.add_table(rows=0, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table2.style = 'Table Grid'
    add_table_row(table2, ["Plan", "Desde", "Ideal para"], header=True)
    add_table_row(table2, ["Básico", "$700.000 COP", "Negocios que inician su presencia digital"])
    add_table_row(table2, ["Profesional", "$1.000.000 COP", "Negocios que quieren destacar y vender más"])
    add_table_row(table2, ["Premium", "$1.800.000 COP", "Empresas que buscan una web completa y optimizada"])

    add_title(doc, "8. Próximos Pasos")
    add_body(doc, "Si esta propuesta te llama la atención, estos son los siguientes pasos:")
    pasos_finales = [
        "✅ Agendemos una llamada rápida por WhatsApp para resolver dudas",
        "✅ Confirmamos el plan y ajustamos detalles",
        "✅ Recibes tu cotización formal y el contrato",
        "✅ Arrancamos con el desarrollo",
    ]
    for pf in pasos_finales:
        add_bullet(doc, pf)

    add_separator(doc)
    add_body(doc, f"Escríbenos por WhatsApp al {WHATSAPP} o al correo {EMAIL}", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_body(doc, f"{COMPANY} — {DOMAIN}", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=LIGHT_GRAY)

    doc.save(os.path.join(OUT_DIR, "05_Propuesta_Comercial.docx"))
    print("[OK] 05_Propuesta_Comercial.docx")


# ─── MAIN ────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generando documentos legales de NuncaCierro...\n")
    create_cotizacion()
    create_contrato()
    create_recibo_caja()
    create_acta_entrega()
    create_propuesta()
    print(f"\nOK - Los 5 documentos fueron generados en:\n   {OUT_DIR}")
